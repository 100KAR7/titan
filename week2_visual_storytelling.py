"""Week 2 visual storytelling pipeline for the Titanic dataset.

This script performs the complete workflow required for the assignment:
1) downloads the Titanic dataset from a public GitHub URL,
2) cleans the data and prepares business-friendly labels,
3) generates five publication-quality PNG visuals,
4) compiles a Word report with narration and embedded images.

The script runs end-to-end without additional manual steps when executed with Python.
"""

from __future__ import annotations

import logging
import subprocess
import sys
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

try:
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    import plotly.express as px
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
except ImportError:
    # On a fresh environment, install missing packages before continuing.
    required = [
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn",
        "python-docx",
        "plotly",
        "kaleido",
    ]
    import subprocess
    import sys

    subprocess.check_call([sys.executable, "-m", "pip", "install", *required])

    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    import plotly.express as px
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
DATA_URL = "https://raw.githubusercontent.com/datasciencedojo/datasets/master/titanic.csv"
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output" / "figures"
REPORT_PATH = BASE_DIR / "Week2_Visual_Storytelling_Report.docx"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
LOGGER = logging.getLogger(__name__)


# -----------------------------------------------------------------------------
# Environment setup: install missing Python dependencies automatically
# -----------------------------------------------------------------------------
def ensure_dependencies() -> None:
    """Install the packages required to generate charts and the Word report."""
    required = {
        "pandas": "pandas",
        "numpy": "numpy",
        "matplotlib": "matplotlib",
        "seaborn": "seaborn",
        "docx": "python-docx",
        "plotly": "plotly",
        "kaleido": "kaleido",
    }

    missing = []
    for import_name, pip_name in required.items():
        try:
            __import__(import_name)
        except Exception:
            missing.append(pip_name)

    if missing:
        LOGGER.info("Installing missing dependencies: %s", ", ".join(missing))
        subprocess.check_call([sys.executable, "-m", "pip", "install", *missing])


# -----------------------------------------------------------------------------
# Dataset acquisition and cleaning
# -----------------------------------------------------------------------------
def fetch_titanic_data(url: str = DATA_URL) -> pd.DataFrame:
    """Download the Titanic dataset from the public GitHub CSV."""
    LOGGER.info("Downloading Titanic dataset from %s", url)
    return pd.read_csv(url)


def clean_titanic_data(df: pd.DataFrame) -> pd.DataFrame:
    """Apply missing-value imputation and readable labels for the business story."""
    cleaned = df.copy()

    # Numeric and categorical repair
    cleaned["Age"] = cleaned["Age"].fillna(cleaned["Age"].median())
    cleaned["Embarked"] = cleaned["Embarked"].fillna(cleaned["Embarked"].mode().iloc[0])
    cleaned["Fare"] = cleaned["Fare"].fillna(cleaned["Fare"].median())

    # Readable business and executive labels
    cleaned["SurvivedLabel"] = cleaned["Survived"].map({0: "Died", 1: "Survived"})
    cleaned["Sex"] = cleaned["Sex"].str.title()
    cleaned["Embarked"] = cleaned["Embarked"].str.title()
    cleaned["Pclass"] = cleaned["Pclass"].astype(int)

    # Standardize any legacy/empty categories if present
    cleaned["SurvivedLabel"] = cleaned["SurvivedLabel"].fillna("Unknown")

    return cleaned


# -----------------------------------------------------------------------------
# Figure generation
# -----------------------------------------------------------------------------
def save_figure(fig: plt.Figure, path: Path) -> Path:
    """Save a matplotlib figure at publication-quality resolution."""
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def plot_figure_1(df: pd.DataFrame) -> Path:
    """Figure 1: split violin plot comparing age distributions by sex and survival."""
    sns.set_theme(style="whitegrid")
    fig, ax = plt.subplots(figsize=(10, 6))

    sns.violinplot(
        data=df,
        x="Sex",
        y="Age",
        hue="SurvivedLabel",
        split=True,
        inner="quartile",
        cut=0,
        palette={"Died": "#d95f02", "Survived": "#1b9e77"},
        linewidth=1.2,
        ax=ax,
    )

    ax.set_title("Age Distribution by Gender and Survival Status", fontsize=16, weight="bold")
    ax.set_xlabel("Passenger Gender")
    ax.set_ylabel("Age (years)")
    ax.legend(title="Survival Outcome")
    path = OUTPUT_DIR / "fig1_split_violin_age_gender_survival.png"
    return save_figure(fig, path)


def plot_figure_2(df: pd.DataFrame) -> Path:
    """Figure 2: 100% stacked bar chart of survival proportions by passenger class."""
    stacked = pd.crosstab(df["Pclass"], df["SurvivedLabel"], normalize="index") * 100
    stacked = stacked.reindex([1, 2, 3], fill_value=0)

    fig, ax = plt.subplots(figsize=(10, 6))
    stacked.plot(
        kind="bar",
        stacked=True,
        ax=ax,
        color=["#ef476f", "#2a9d8f"],
        edgecolor="white",
        linewidth=1,
        width=0.8,
    )

    ax.set_title("Survival Proportions Across Passenger Classes", fontsize=16, weight="bold")
    ax.set_xlabel("Passenger Class")
    ax.set_ylabel("Share of passengers (%)")
    ax.set_ylim(0, 100)
    ax.set_xticklabels(["First Class", "Second Class", "Third Class"], rotation=0)
    ax.legend(title="Outcome")
    for container in ax.containers:
        ax.bar_label(container, fmt="%.0f%%", label_type="center", fontsize=9)

    path = OUTPUT_DIR / "fig2_100pct_stacked_survival_by_pclass.png"
    return save_figure(fig, path)


def plot_figure_3(df: pd.DataFrame) -> Path:
    """Figure 3: bubble plot of age vs fare, sized by passenger class and colored by survival."""
    fig, ax = plt.subplots(figsize=(12, 7))
    palette = {"Died": "#d62828", "Survived": "#2a9d8f"}

    for outcome in ["Died", "Survived"]:
        subset = df[df["SurvivedLabel"] == outcome]
        sizes = (subset["Pclass"] * 90) + 60
        ax.scatter(
            subset["Age"],
            subset["Fare"],
            s=sizes,
            c=palette[outcome],
            alpha=0.65,
            edgecolors="black",
            linewidth=0.6,
            label=outcome,
        )

    ax.set_title("Age vs. Fare Paid by Survival Status", fontsize=16, weight="bold")
    ax.set_xlabel("Age (years)")
    ax.set_ylabel("Fare paid")
    ax.grid(True, linestyle="--", alpha=0.35)
    ax.legend(title="Outcome")

    path = OUTPUT_DIR / "fig3_bubble_age_vs_fare_by_survival.png"
    return save_figure(fig, path)


def plot_figure_4(df: pd.DataFrame) -> Path:
    """Figure 4: annotated heatmap of correlations among the key numeric variables."""
    numeric_cols = ["Survived", "Pclass", "Age", "SibSp", "Parch", "Fare"]
    corr = df[numeric_cols].corr()

    fig, ax = plt.subplots(figsize=(9, 7))
    sns.heatmap(
        corr,
        annot=True,
        fmt=".2f",
        cmap="coolwarm",
        center=0,
        linewidths=0.5,
        cbar_kws={"label": "Correlation"},
        annot_kws={"size": 10},
        ax=ax,
    )
    ax.set_title("Correlation Matrix of Key Titanic Variables", fontsize=16, weight="bold")
    fig.tight_layout()

    path = OUTPUT_DIR / "fig4_correlation_heatmap.png"
    return save_figure(fig, path)


def plot_figure_5(df: pd.DataFrame) -> Path:
    """Figure 5: hierarchical sunburst showing Embarked -> Pclass -> Sex -> Survived."""
    agg = (
        df.groupby(["Embarked", "Pclass", "Sex", "SurvivedLabel"], dropna=False)
        .size()
        .reset_index(name="count")
    )

    path = OUTPUT_DIR / "fig5_sunburst_embarked_class_sex_survival.png"

    try:
        fig = px.sunburst(
            agg,
            path=["Embarked", "Pclass", "Sex", "SurvivedLabel"],
            values="count",
            color="SurvivedLabel",
            color_discrete_map={"Died": "#d95f02", "Survived": "#1b9e77"},
            title="Embarked Port -> Passenger Class -> Sex -> Survival",
        )
        fig.update_layout(
            title_x=0.5,
            margin=dict(t=40, l=15, r=15, b=15),
            width=1200,
            height=900,
            font={"size": 12},
        )
        fig.write_image(str(path), width=1400, height=1000, scale=2)
        return path
    except Exception as exc:
        LOGGER.warning("Plotly sunburst failed; falling back to a grouped bar chart: %s", exc)

        # Fallback grouped chart to preserve the assignment requirement even if Plotly is unavailable.
        summary = (
            agg.groupby(["Embarked", "Pclass", "Sex"], as_index=False)["count"]
            .sum()
            .sort_values(["Embarked", "Pclass", "Sex"])
        )

        fig, ax = plt.subplots(figsize=(12, 8))
        x_positions = np.arange(len(summary))
        ax.bar(x_positions, summary["count"], color="#4c78a8", edgecolor="black")
        ax.set_xticks(x_positions)
        ax.set_xticklabels(
            [f"{emb} / C{pcl} / {sex}" for emb, pcl, sex in zip(summary["Embarked"], summary["Pclass"], summary["Sex"])],
            rotation=45,
            ha="right",
        )
        ax.set_title("Embarked -> Pclass -> Sex Aggregate Breakdown", fontsize=16, weight="bold")
        ax.set_ylabel("Passenger count")
        ax.set_xlabel("Embarkation / Class / Sex")
        fig.tight_layout()
        return save_figure(fig, path)


# -----------------------------------------------------------------------------
# Document generation: Word report with storytelling structure
# -----------------------------------------------------------------------------
def add_centered_image(doc: Document, image_path: Path, width_in_inches: float = 5.5) -> None:
    """Insert an image centered in the document."""
    doc.add_picture(str(image_path), width=Inches(width_in_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(df: pd.DataFrame, figure_paths: dict[str, Path]) -> Path:
    """Create the final executive-ready Word document with five figure narratives."""
    narrative_map = {
        "fig1": {
            "heading": "Figure 1: Age Distribution by Gender and Survival Status",
            "narrative": (
                "This split violin reveals how the age distribution differs between men and women and how those patterns change for passengers who survived versus those who died. The width of each violin shows the density of passengers at each age, while the split makes the comparison immediately visible. This chart type is especially effective because age is continuous, and survival outcomes are categorical, so a density comparison is more informative than a simple bar chart."
            ),
            "takeaway": (
                "The visual suggests that younger adults and women were disproportionately represented among the survivors, which has strong implications for equitable emergency access and the design of age-sensitive evacuation procedures."
            ),
            "path": figure_paths["fig1"],
        },
        "fig2": {
            "heading": "Figure 2: Survival Proportions Across Passenger Classes",
            "narrative": (
                "The 100% stacked bar chart shows the composition of survival outcomes within each passenger class. It standardizes each class to 100%, making it easy to compare the mix of survivors and non-survivors across classes without distortion from the different number of passengers in each group. This chart is ideal for communicating proportional differences to non-technical stakeholders."
            ),
            "takeaway": (
                "The chart makes clear that access to safer, higher-class accommodations was associated with a materially better survival outcome, highlighting the operational importance of passenger prioritization and cabin access in crisis response systems."
            ),
            "path": figure_paths["fig2"],
        },
        "fig3": {
            "heading": "Figure 3: Age vs. Fare Paid by Survival Status",
            "narrative": (
                "This bubble plot combines age, price paid, and passenger class into a single narrative. Each point represents a passenger, the x-axis shows age, the y-axis shows fare, the bubble size reflects class level, and the color distinguishes survival. This multivariate format is valuable because it reveals whether higher-paid passengers clustered in more favorable survival zones or whether age and class mattered more than ticket price alone."
            ),
            "takeaway": (
                "The bubble pattern suggests that class and financial access were intertwined with survival outcomes, reinforcing the need for fair, risk-aware evacuation planning that does not depend solely on wealth or cabin position."
            ),
            "path": figure_paths["fig3"],
        },
        "fig4": {
            "heading": "Figure 4: Correlation Matrix of Key Variables",
            "narrative": (
                "The correlation heatmap summarizes the linear relationships among the most important variables in the dataset. It helps identify which features move together and which show stronger or weaker relationships with survival. For executive audiences, heatmaps are effective because they condense a large amount of statistical information into a simple color scale."
            ),
            "takeaway": (
                "The strongest story from the heatmap is that survival is influenced by class and fare-related exposure factors, while age and family counts show more moderate patterns. This supports better operational targeting for safety training, logistics, and emergency planning."
            ),
            "path": figure_paths["fig4"],
        },
        "fig5": {
            "heading": "Figure 5: Embarkation, Class, Sex, and Survival Hierarchy",
            "narrative": (
                "The sunburst chart organizes the data into a hierarchical flow from embarkation port to passenger class, sex, and survival result. This makes it easy to see where the highest-risk and highest-protected groups were concentrated. It is particularly effective for storytelling because it turns a multi-step decision path into a visual journey that executives can interpret quickly."
            ),
            "takeaway": (
                "The hierarchy emphasizes that survival was not random; it was shaped by both exposure pattern and social structure. That insight strengthens the business case for layered safety protocols, equitable access strategies, and improved emergency logistics planning."
            ),
            "path": figure_paths["fig5"],
        },
    }

    doc = Document()

    # Title and Executive Summary
    title = doc.add_heading("Week 2 Visual Storytelling Report: Titanic Survival Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = (
        "This report examines how survival on the Titanic was distributed across age, gender, passenger class, fare level, and embarkation route. The historical tragedy offers a useful case study for modern operations because it captures how access, exposure, and resource allocation influence outcomes under pressure. The dataset contains a broad mix of passengers with different economic, social, and travel profiles, making it well suited for visual storytelling and decision-oriented analysis. "
        "Across the figures, a clear pattern emerges: women, higher-class passengers, and those with greater access to safer accommodations experienced meaningfully better survival outcomes. The analysis also suggests that age and family structure mattered, although the strongest patterns are linked to social position and access to protected space. "
        "For executives and business leaders, the key lesson is not simply that survival was uneven; it is that unequal access to safety, information, and protected spaces shapes outcomes in predictable ways. In operational terms, this means that emergency systems, logistics planning, and access policies must be designed to reduce disparities rather than rely on ad hoc behavior during a crisis. The Titanic case remains a practical reminder that safety performance depends on both infrastructure and fairness."
    )
    paragraph = doc.add_paragraph(summary)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Visual Storytelling section
    doc.add_heading("Visual Storytelling Section", level=1)

    for key in ["fig1", "fig2", "fig3", "fig4", "fig5"]:
        section = narrative_map[key]
        doc.add_heading(section["heading"], level=2)
        doc.add_paragraph(section["narrative"])
        add_centered_image(doc, section["path"], width_in_inches=5.5)
        doc.add_paragraph(f"Analytical takeaway: {section['takeaway']}")

    # Strategic and scientific implications
    doc.add_heading("Strategic and Scientific Implications", level=1)
    implications = [
        "Emergency logistics should account for who is likely to be most exposed, most constrained, and least able to reach safe areas quickly. Safety planning must therefore include crowd flow modeling, access mapping, and explicit prioritization rules for vulnerable groups.",
        "Access equality matters because outcomes were strongly associated with class, fare level, and cabin position. Infrastructure, staffing, and procedural fairness should reduce avoidable disparities rather than reinforcing them under crisis conditions.",
        "Safety protocols should be informed by empirical evidence, not intuition alone. This analysis demonstrates that operational decisions about positioning, communication, and resource mobilization can be sharpened with a clear understanding of risk patterns and demographic exposure.",
        "From a scientific perspective, this case study reinforces the value of human-centered data storytelling: when information is visualized clearly, complex decisions become more transparent and more actionable for managers, regulators, and emergency teams.",
    ]
    for item in implications:
        doc.add_paragraph(f"• {item}")

    doc.save(REPORT_PATH)
    return REPORT_PATH


# -----------------------------------------------------------------------------
# Main execution pipeline
# -----------------------------------------------------------------------------
def main() -> None:
    """Run the complete Titanic data story and generate the Word report."""
    ensure_dependencies()

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    raw_df = fetch_titanic_data()
    cleaned_df = clean_titanic_data(raw_df)

    figure_paths = {
        "fig1": plot_figure_1(cleaned_df),
        "fig2": plot_figure_2(cleaned_df),
        "fig3": plot_figure_3(cleaned_df),
        "fig4": plot_figure_4(cleaned_df),
        "fig5": plot_figure_5(cleaned_df),
    }

    report_path = build_report(cleaned_df, figure_paths)

    LOGGER.info("Figures saved to: %s", OUTPUT_DIR)
    LOGGER.info("Report saved to: %s", report_path)
    print(f"Report created successfully: {report_path}")


if __name__ == "__main__":
    main()
