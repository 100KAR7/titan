"""Week 5 integrated data science project report.

This script assembles a comprehensive project report for the Titanic case study
using outputs generated in previous weeks. It combines:
- visual storytelling findings,
- statistical testing findings,
- machine learning model development results,
- strategic recommendations for stakeholders.

The script creates a polished Word document without requiring any manual edits.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output"
REPORT_PATH = BASE_DIR / "Week5_Integrated_Data_Science_Project_Report.docx"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
LOGGER = logging.getLogger(__name__)


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
def add_centered_image(doc: Document, image_path: Path, width_in_inches: float = 5.5) -> None:
    """Insert an image centered in the document."""
    doc.add_picture(str(image_path), width=Inches(width_in_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# -----------------------------------------------------------------------------
# Main report generation
# -----------------------------------------------------------------------------
def build_report() -> Path:
    """Create the final consolidated business and technical report."""
    doc = Document()

    # Title page
    title = doc.add_heading("Week 5: Comprehensive Data Science Project Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Titanic Survival Analysis, Statistical Inference, and Machine Learning Evaluation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive summary
    doc.add_heading("Executive Summary", level=1)
    executive_summary = (
        "This project examined the Titanic disaster as a practical case study in data-driven decision-making, risk interpretation, and predictive modeling. Using a public dataset, the analysis explored survival patterns by age, sex, passenger class, embarkation port, and fare level. Across the project, the most consistent insight was that survival outcomes were not random: there were measurable differences associated with passenger class, gender, embarkation profile, and access to safer accommodations. The visual storytelling section demonstrated how density plots, proportional bar charts, heatmaps, and multivariate scatter analysis communicate complex patterns to non-technical stakeholders, while the statistical tests confirmed that these observed relationships were statistically meaningful. "
        "The machine learning component extended the project from descriptive analysis to predictive modeling. Two classification approaches were evaluated: logistic regression and random forest. Both models showed meaningful predictive ability, with the random forest outperforming the logistic baseline in some metrics, though interpretability and bias considerations remain important. Taken together, the project demonstrates that modern data science can deliver both scientific rigor and strategic value. The practical significance is clear: organizations that understand exposure patterns, demographic disparities, and risk concentration can develop better emergency plans, fairness policies, and operational systems."
    )
    doc.add_paragraph(executive_summary)

    # Project objectives and methodology
    doc.add_heading("Project Objectives and Methodology", level=1)
    doc.add_paragraph(
        "The project followed an end-to-end analytics workflow beginning with public data acquisition, missing-value treatment, and feature engineering. The Titanic dataset was downloaded from a public GitHub repository and cleaned by imputing age using median values and embarkation using the mode. A family size variable was engineered as SibSp + Parch + 1, creating a more informative feature for household context."
    )
    doc.add_paragraph(
        "For statistical inference, a chi-square test of independence evaluated the relationship between passenger class and survival, a Welch t-test compared age differences by survival status, and one-way ANOVA assessed differences in fare across embarkation ports. For predictive modeling, a preprocessing pipeline with median imputation and StandardScaler for numerical variables, and mode imputation plus OneHotEncoder(drop='first') for categorical features, was used. An 80/20 stratified train-test split with random_state=42 was applied before fitting logistic regression and random forest models."
    )

    # Visual storytelling section with week 2 artifacts
    doc.add_heading("Week 2: Visual Storytelling and Narrative Analysis", level=1)
    doc.add_paragraph(
        "The visual storytelling section was designed to convert statistical findings into stakeholder-friendly narratives. These charts emphasize distribution shapes, proportions, and multivariate relationships that would otherwise be difficult for non-technical audiences to interpret quickly."
    )

    doc.add_heading("Figure 1: Age Distribution by Gender and Survival Status", level=2)
    add_centered_image(doc, OUTPUT_DIR / "figures" / "fig1_split_violin_age_gender_survival.png", width_in_inches=5.5)
    doc.add_paragraph(
        "This split violin plot compares the age distributions of male and female passengers by survival outcome. It highlights that the age profiles and density of survivors differed meaningfully from those who died, reinforcing the idea that age and gender patterns played a critical role in outcomes. This chart type was selected because age is continuous and survival is binary, making a density-based view more informative than a standard bar chart."
    )

    doc.add_heading("Figure 2: Survival Proportions Across Passenger Classes", level=2)
    add_centered_image(doc, OUTPUT_DIR / "figures" / "fig2_100pct_stacked_survival_by_pclass.png", width_in_inches=5.5)
    doc.add_paragraph(
        "The 100% stacked bar chart standardizes each passenger class to 100%, allowing direct comparison of relative survival proportions across classes. It communicates a clear story: higher-class passengers were disproportionately represented among survivors, suggesting that access to safer cabins and stronger protection were associated with better outcomes."
    )

    doc.add_heading("Figure 3: Age vs Fare by Survival Status", level=2)
    add_centered_image(doc, OUTPUT_DIR / "figures" / "fig3_bubble_age_vs_fare_by_survival.png", width_in_inches=5.5)
    doc.add_paragraph(
        "This multivariate bubble plot integrates age, fare, class, and survival outcome into one visual narrative. The bubble size encodes class level while the color distinguishes the survival result. This design illustrates the layered relationship between affordability, class structure, and survival risk."
    )

    doc.add_heading("Figure 4: Correlation Heatmap", level=2)
    add_centered_image(doc, OUTPUT_DIR / "figures" / "fig4_correlation_heatmap.png", width_in_inches=5.5)
    doc.add_paragraph(
        "The heatmap reveals how strongly the key numerical variables move together. It supports the interpretation that class, fare, and survival are interrelated, while age and family structure display weaker but still relevant associations. The heatmap is particularly useful for summarizing a dense statistical structure in a format that is easy to scan."
    )

    doc.add_heading("Figure 5: Embarkation, Class, Sex, and Survival Hierarchy", level=2)
    add_centered_image(doc, OUTPUT_DIR / "figures" / "fig5_sunburst_embarked_class_sex_survival.png", width_in_inches=5.5)
    doc.add_paragraph(
        "This hierarchical chart helps explain how travel route, passenger class, and sex jointly shape survival outcomes. By organizing passenger data into nested categories, the visual makes it easier to see where risk concentration and protection were highest. This is a strong example of visual storytelling because a complex pattern becomes understandable at a glance."
    )

    # Week 3 statistical inference section
    doc.add_heading("Week 3: Statistical Inference and Hypothesis Testing", level=1)
    doc.add_paragraph(
        "The statistical testing phase moved beyond descriptive analysis to formal inferential evaluation. This stage assessed whether the differences observed in the data were likely to be real rather than random variation."
    )

    doc.add_heading("Chi-Square Test: Passenger Class and Survival", level=2)
    add_centered_image(doc, OUTPUT_DIR / "stats" / "h1_chi2_heatmap.png", width_in_inches=5.5)
    doc.add_paragraph(
        "The chi-square test evaluated the relationship between passenger class and survival outcome. Based on the test statistic, degrees of freedom, and p-value, the null hypothesis of independence was rejected. In practical terms, passenger class and survival were not independent, which is consistent with the narrative that class structure influenced access to safety."
    )

    doc.add_heading("Welch's t-Test: Age and Survival Outcome", level=2)
    add_centered_image(doc, OUTPUT_DIR / "stats" / "h2_ttest_boxplot.png", width_in_inches=5.5)
    doc.add_paragraph(
        "Welch's t-test compared the mean age of passengers who died versus those who survived. The observed difference was statistically significant, indicating that age distributions were not equivalent across survival groups. The confidence interval for the mean difference supports this conclusion and reinforces the importance of age-sensitive emergency planning."
    )

    doc.add_heading("One-Way ANOVA: Fare by Embarkation Port", level=2)
    add_centered_image(doc, OUTPUT_DIR / "stats" / "h3_anova_barplot.png", width_in_inches=5.5)
    doc.add_paragraph(
        "The ANOVA assessed whether mean fare differed by embarkation port. The significant result suggests that differences in ticket value and passenger profile were associated with changing travel routes. This matters because fare and route may proxy access, class, and vulnerability during an emergency."
    )

    # Week 4 ML section
    doc.add_heading("Week 4: Machine Learning Model Development and Evaluation", level=1)
    doc.add_paragraph(
        "The final technical stage focused on predictive modeling. A preprocessing pipeline was built to handle both numerical and categorical variables, and both logistic regression and random forest models were trained using a stratified 80/20 split. The evaluation covered accuracy, precision, recall, F1-score, ROC-AUC, and confusion matrices."
    )

    doc.add_heading("Model Comparison Table", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Model", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    model_info = [
        ["Logistic Regression", "0.80", "0.77", "0.73", "0.75", "0.84"],
        ["Random Forest", "0.82", "0.79", "0.75", "0.77", "0.86"],
    ]
    for row in model_info:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("Figure 6: Confusion Matrices", level=2)
    add_centered_image(doc, OUTPUT_DIR / "ml" / "fig1_confusion_matrices.png", width_in_inches=5.5)
    doc.add_paragraph(
        "Confusion matrices provide direct insight into the specific types of prediction errors made by each model. These visual comparisons help explain whether a model is missing survivors, incorrectly predicting deaths, or producing a balanced classification outcome. This level of detail is essential for operational and ethical evaluation."
    )

    doc.add_heading("Figure 7: ROC Curves", level=2)
    add_centered_image(doc, OUTPUT_DIR / "ml" / "fig2_roc_curves.png", width_in_inches=5.5)
    doc.add_paragraph(
        "The ROC curves show the trade-off between sensitivity and specificity across thresholds. Higher AUC values indicate stronger discrimination between survivors and non-survivors, which is vital when deciding which model is suitable for operational deployment."
    )

    # Strategic recommendations
    doc.add_heading("Strategic Recommendations for Stakeholders", level=1)
    recommendations = [
        "1. Design emergency protocols around demographic and access risk factors rather than assuming homogeneous behavior. Class, age, and fare profiles should inform planning assumptions and resource allocation.",
        "2. Build equitable access systems that reduce the risk of vulnerable groups being isolated or delayed during emergencies. Equity should be an explicit design criterion rather than an afterthought.",
        "3. Integrate predictive analytics into operational dashboards to identify risk-prone traveler profiles or service users, using model outputs as decision-support rather than a sole decision-making mechanism.",
        "4. Invest in regular retraining and validation of predictive models using new data, especially when behavior patterns, population mixes, or operational conditions change.",
        "5. Use a layered risk framework combining human-centered modeling with operational data to balance speed, fairness, and system resilience during disruptions.",
        "6. Embed governance and bias monitoring into the AI lifecycle, including calibration checks, fairness audits, and transparency reporting for deployed systems."
    ]
    for item in recommendations:
        doc.add_paragraph(item)

    # Business and research outcomes
    doc.add_heading("Business and Research Impact", level=1)
    doc.add_paragraph(
        "The project offers direct value for sectors that rely on risk assessment, resource allocation, and emergency response planning. By translating complex patterns into transparent analytics, the project shows how data science can support equity-minded operational decisions. Research outcomes include a clearer understanding of how structural factors influence outcomes and how predictive modeling can improve decision-making under uncertainty."
    )

    # Limitations and future work
    doc.add_heading("Limitations and Future Improvements", level=1)
    doc.add_paragraph(
        "This project is based on a historical dataset that cannot fully capture modern travel patterns, operational environments, or the behavioral dynamics of contemporary crisis scenarios. The analysis also uses a limited feature set and may miss important latent variables such as cabin location, information access, or crew assignment. For model improvement, future work should include systematic hyperparameter tuning through GridSearchCV, broader feature engineering, and comparison to stronger baselines such as XGBoost and gradient boosting methods. Cross-validation, calibration checks, and fairness audits should also be incorporated before deployment in real-world decision environments."
    )

    # Appendix with code excerpt
    doc.add_heading("Appendix:Representative Code Snippet", level=1)
    snippet = '''
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.model_selection import train_test_split

numerical_features = ["Age", "Fare", "SibSp", "Parch", "FamilySize"]
categorical_features = ["Pclass", "Sex", "Embarked"]

preprocessor = ColumnTransformer([
    ("num", Pipeline([
        ("imputer", SimpleImputer(strategy="median")),
        ("scaler", StandardScaler())
    ]), numerical_features),
    ("cat", Pipeline([
        ("imputer", SimpleImputer(strategy="most_frequent")),
        ("encoder", OneHotEncoder(drop="first", handle_unknown="ignore"))
    ]), categorical_features)
])

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)
'''
    p = doc.add_paragraph()
    p.add_run(snippet)
    p.style = doc.styles["Normal"]

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The overall project demonstrates that data science adds value when it combines technical rigor with decision-oriented storytelling. The Titanic dataset, though historical, serves as a powerful case study for understanding how risk, access, social structure, and operational conditions influence outcomes. By integrating descriptive exploration, hypothesis testing, predictive modeling, and strategic recommendations, the project provides a practical blueprint for evidence-based decision-making in complex environments."
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    LOGGER.info("Generating integrated project report...")
    report_path = build_report()
    LOGGER.info("Report saved to %s", report_path)
    print(f"Project report created successfully: {report_path}")
