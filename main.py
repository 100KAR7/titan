"""End-to-end Titanic data analysis and Word report generation.

Run with::

    python main.py

The script downloads the source data, cleans it, creates three PNG figures,
and writes Data_Analysis_Report.docx plus supporting artifacts to output/.
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DATA_URL = "https://raw.githubusercontent.com/datasciencedojo/datasets/master/titanic.csv"
PROJECT_ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_ROOT / "output"
FIGURES_DIR = OUTPUT_DIR / "figures"
REPORT_PATH = OUTPUT_DIR / "Data_Analysis_Report.docx"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
LOGGER = logging.getLogger(__name__)


def acquire_data(url: str = DATA_URL) -> pd.DataFrame:
    """Download the Titanic CSV into a DataFrame."""
    LOGGER.info("Downloading dataset from %s", url)
    return pd.read_csv(url)


def clean_data(data: pd.DataFrame) -> tuple[pd.DataFrame, dict[str, Any]]:
    """Apply documented missing-value, duplicate, and dtype transformations."""
    cleaned = data.copy()
    missing_before = cleaned.isna().sum()
    duplicate_rows = int(cleaned.duplicated().sum())
    cabin_missing_rate = float(cleaned["Cabin"].isna().mean()) if "Cabin" in cleaned else 0.0

    if "Age" in cleaned:
        cleaned["Age"] = cleaned["Age"].fillna(cleaned["Age"].median())
    if "Cabin" in cleaned and cabin_missing_rate > 0.70:
        cleaned = cleaned.drop(columns="Cabin")
    if "Embarked" in cleaned:
        cleaned["Embarked"] = cleaned["Embarked"].fillna(cleaned["Embarked"].mode().iloc[0])
    cleaned = cleaned.drop_duplicates().reset_index(drop=True)

    for column in ("Survived", "Pclass"):
        if column in cleaned:
            cleaned[column] = cleaned[column].astype("category")
    for column in ("Sex", "Embarked"):
        if column in cleaned:
            cleaned[column] = cleaned[column].astype("category")

    verification = {
        "rows_before": len(data),
        "rows_after": len(cleaned),
        "columns_before": len(data.columns),
        "columns_after": len(cleaned.columns),
        "duplicates_removed": duplicate_rows,
        "cabin_missing_rate": cabin_missing_rate,
        "missing_before": missing_before[missing_before > 0].to_dict(),
        "missing_after": cleaned.isna().sum()[cleaned.isna().sum() > 0].to_dict(),
        "dtypes": {column: str(dtype) for column, dtype in cleaned.dtypes.items()},
    }
    return cleaned, verification


def summarize_data(data: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Return numerical and categorical summary statistics."""
    numerical = data.select_dtypes(include=np.number).describe().round(2).T
    categorical_columns = data.select_dtypes(include=["str", "category"]).columns
    categorical = data[categorical_columns].describe().T if len(categorical_columns) else pd.DataFrame()
    return numerical, categorical


def create_visualizations(data: pd.DataFrame) -> dict[str, Path]:
    """Create and save publication-quality EDA charts."""
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)
    sns.set_theme(style="whitegrid", context="notebook", palette="deep")
    paths: dict[str, Path] = {}

    figure, axis = plt.subplots(figsize=(8, 5), constrained_layout=True)
    sns.histplot(data=data, x="Age", bins=30, kde=True, color="#2a6f97", ax=axis)
    axis.set(title="Age Distribution", xlabel="Age (years)", ylabel="Passenger count")
    paths["Age Distribution"] = FIGURES_DIR / "age_distribution.png"
    figure.savefig(paths["Age Distribution"], dpi=300, bbox_inches="tight")
    plt.close(figure)

    figure, axis = plt.subplots(figsize=(8, 5), constrained_layout=True)
    survival_data = data.assign(Survived=data["Survived"].astype(str))
    sns.countplot(data=survival_data, x="Sex", hue="Survived", ax=axis)
    axis.set(title="Survival Count by Sex", xlabel="Sex", ylabel="Passenger count")
    axis.legend(title="Survived", labels=["No", "Yes"])
    paths["Survival Count by Sex"] = FIGURES_DIR / "survival_count_by_sex.png"
    figure.savefig(paths["Survival Count by Sex"], dpi=300, bbox_inches="tight")
    plt.close(figure)

    figure, axis = plt.subplots(figsize=(8, 5), constrained_layout=True)
    fare_data = data.assign(Pclass=data["Pclass"].astype(str))
    sns.boxplot(data=fare_data, x="Pclass", y="Fare", hue="Pclass", legend=False, ax=axis)
    axis.set(title="Fare Distribution across Passenger Classes", xlabel="Passenger class", ylabel="Fare")
    paths["Fare Distribution across Passenger Classes"] = FIGURES_DIR / "fare_distribution_by_class.png"
    figure.savefig(paths["Fare Distribution across Passenger Classes"], dpi=300, bbox_inches="tight")
    plt.close(figure)
    return paths


def source_for(function: Any) -> str:
    """Get a readable source snippet for inclusion in the report."""
    import inspect

    try:
        return inspect.getsource(function)
    except (OSError, TypeError):
        return "Source unavailable in this execution environment."


def set_cell_shading(cell: Any, fill: str) -> None:
    """Apply a background color to a table cell."""
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_dataframe_table(document: Document, frame: pd.DataFrame, max_rows: int = 12) -> None:
    """Add a compact DataFrame table to the report."""
    if frame.empty:
        document.add_paragraph("No variables were available for this summary.")
        return
    display_frame = frame.reset_index().head(max_rows)
    table = document.add_table(rows=1, cols=len(display_frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Shading Accent 1"
    for index, column in enumerate(display_frame.columns):
        table.rows[0].cells[index].text = str(column)
        set_cell_shading(table.rows[0].cells[index], "1F4E79")
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    for _, row in display_frame.iterrows():
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)


def add_code_block(document: Document, code: str) -> None:
    """Add monospaced code with a subtle shaded background."""
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["No Spacing"]
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    for line_number, line in enumerate(code.strip().splitlines(), start=1):
        run = paragraph.add_run(f"{line_number:>3}  {line}\n")
        run.font.name = "Consolas"
        run.font.size = Pt(8)


def add_figure(document: Document, title: str, path: Path, insight: str) -> None:
    """Embed a chart with a caption and interpretation."""
    document.add_picture(str(path), width=Inches(6.2))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.add_run(f"Figure: {title}").bold = True
    document.add_paragraph(f"Analytical insight: {insight}")


def build_report(
    raw_data: pd.DataFrame,
    cleaned_data: pd.DataFrame,
    verification: dict[str, Any],
    numerical_summary: pd.DataFrame,
    categorical_summary: pd.DataFrame,
    figures: dict[str, Path],
) -> Path:
    """Build and save the formatted analysis report."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    document.core_properties.title = "Comprehensive Data Analysis Report: Titanic Survival Dataset"

    title = document.add_heading("Comprehensive Data Analysis Report: Titanic Survival Dataset", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Prepared from {len(raw_data):,} source rows. Generated automatically with pandas, "
        "NumPy, Seaborn, Matplotlib, and python-docx."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Section 1: Methodology & Data Acquisition", 1)
    document.add_paragraph(
        "The pipeline retrieves the public Titanic CSV directly from GitHub, preserves the raw "
        "DataFrame for auditability, applies deterministic preprocessing, calculates descriptive "
        "statistics, creates reusable visualizations, and packages the results into this report."
    )
    document.add_heading("Acquisition and analysis code", 2)
    add_code_block(document, source_for(acquire_data) + "\n\n" + source_for(summarize_data))

    document.add_heading("Section 2: Data Cleaning & Preprocessing", 1)
    document.add_paragraph(
        "Age is imputed with its median because it is numeric and the median is robust to skew and "
        "outliers. Cabin is removed because more than 70% of its values are missing, making reliable "
        "single-value imputation inappropriate. Missing Embarked values are filled with the mode, "
        "the most frequent observed embarkation port. Duplicate rows are removed, and key discrete "
        "variables are converted to categorical dtype for correct analysis semantics."
    )
    document.add_heading("Cleaning code", 2)
    add_code_block(document, source_for(clean_data))
    document.add_heading("Verification output", 2)
    verification_lines = [
        f"Rows: {verification['rows_before']} -> {verification['rows_after']}",
        f"Columns: {verification['columns_before']} -> {verification['columns_after']}",
        f"Duplicate rows removed: {verification['duplicates_removed']}",
        f"Cabin missing rate before cleaning: {verification['cabin_missing_rate']:.1%}",
        f"Missing values before: {verification['missing_before']}",
        f"Missing values after: {verification['missing_after'] or 'none'}",
        f"Dtypes after cleaning: {verification['dtypes']}",
    ]
    document.add_paragraph("\n".join(verification_lines))
    document.add_paragraph(f"Final dataset dimensions: {cleaned_data.shape[0]} rows x {cleaned_data.shape[1]} columns.")

    document.add_heading("Section 3: Exploratory Data Analysis & Visualizations", 1)
    document.add_heading("Numerical summary statistics", 2)
    add_dataframe_table(document, numerical_summary)
    document.add_heading("Categorical summary statistics", 2)
    add_dataframe_table(document, categorical_summary)
    add_figure(
        document,
        "Age Distribution",
        figures["Age Distribution"],
        "The distribution is concentrated among young and middle-aged adults, with the KDE showing the overall shape after median imputation.",
    )
    add_figure(
        document,
        "Survival Count by Sex",
        figures["Survival Count by Sex"],
        "Female passengers have substantially more survivors and fewer deaths than male passengers, revealing a strong association between sex and survival outcome.",
    )
    add_figure(
        document,
        "Fare Distribution across Passenger Classes",
        figures["Fare Distribution across Passenger Classes"],
        "Fare generally increases with passenger class, while the spread and high-value outliers show that price varied considerably within classes.",
    )

    document.add_heading("Section 4: Summary of Key Insights & Future Scope", 1)
    document.add_paragraph(
        "The analysis indicates that passenger sex and class are important descriptive correlates of "
        "Titanic survival, while passenger ages are broadly distributed across the adult population. "
        "Future work could add survival-rate confidence intervals, investigate interactions among sex, "
        "class, age, and family size, and evaluate predictive models with cross-validation. Any model "
        "should use a documented train/test split and avoid leakage from imputation or feature engineering."
    )
    document.add_paragraph("Artifacts are stored alongside this report in the output directory.")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_PATH)
    return REPORT_PATH


def main() -> None:
    """Run the complete data analysis workflow."""
    os.environ.setdefault("MPLBACKEND", "Agg")
    raw_data = acquire_data()
    cleaned_data, verification = clean_data(raw_data)
    numerical_summary, categorical_summary = summarize_data(cleaned_data)
    figures = create_visualizations(cleaned_data)
    report_path = build_report(
        raw_data,
        cleaned_data,
        verification,
        numerical_summary,
        categorical_summary,
        figures,
    )
    LOGGER.info("Saved %d figures to %s", len(figures), FIGURES_DIR)
    LOGGER.info("Saved report to %s", report_path)


if __name__ == "__main__":
    main()