"""Week 4 machine learning model development pipeline for the Titanic dataset.

This script performs the complete workflow required for the assignment:
1) downloads the Titanic dataset,
2) engineers features and builds a preprocessing pipeline,
3) trains two classification models,
4) evaluates them with standard metrics,
5) saves publication-quality figures,
6) compiles a Word report summarizing the findings.

The script is designed to run end-to-end without manual intervention in a fresh environment.
"""

from __future__ import annotations

import logging
import os
import subprocess
import sys
from pathlib import Path

# Ensure a non-interactive plotting backend for headless execution.
os.environ.setdefault("MPLBACKEND", "Agg")

# Graceful dependency bootstrap for fresh environments.
try:
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from sklearn.compose import ColumnTransformer
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.impute import SimpleImputer
    from sklearn.linear_model import LogisticRegression
    from sklearn.metrics import (
        accuracy_score,
        confusion_matrix,
        f1_score,
        precision_score,
        recall_score,
        roc_auc_score,
        roc_curve,
    )
    from sklearn.model_selection import train_test_split
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import OneHotEncoder, StandardScaler
except ImportError:
    required = [
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn",
        "python-docx",
        "scikit-learn",
    ]
    subprocess.check_call([sys.executable, "-m", "pip", "install", *required])

    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from sklearn.compose import ColumnTransformer
    from sklearn.ensemble import RandomForestClassifier
    from sklearn.impute import SimpleImputer
    from sklearn.linear_model import LogisticRegression
    from sklearn.metrics import (
        accuracy_score,
        confusion_matrix,
        f1_score,
        precision_score,
        recall_score,
        roc_auc_score,
        roc_curve,
    )
    from sklearn.model_selection import train_test_split
    from sklearn.pipeline import Pipeline
    from sklearn.preprocessing import OneHotEncoder, StandardScaler

# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
DATA_URL = "https://raw.githubusercontent.com/datasciencedojo/datasets/master/titanic.csv"
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output" / "ml"
REPORT_PATH = BASE_DIR / "Week4_Machine_Learning_Model_Development_&_Evaluation.docx"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
LOGGER = logging.getLogger(__name__)


# -----------------------------------------------------------------------------
# Data acquisition and preprocessing
# -----------------------------------------------------------------------------
def fetch_data(url: str = DATA_URL) -> pd.DataFrame:
    """Download the Titanic dataset from the provided public GitHub CSV."""
    LOGGER.info("Downloading Titanic dataset from %s", url)
    return pd.read_csv(url)


def prepare_dataset(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.Series]:
    """Create the target and engineered features for model training."""
    data = df.copy()

    # Fill missing values as specified.
    if "Age" in data.columns:
        data["Age"] = data["Age"].fillna(data["Age"].median())
    if "Embarked" in data.columns:
        data["Embarked"] = data["Embarked"].fillna(data["Embarked"].mode().iloc[0])
    if "Fare" in data.columns:
        data["Fare"] = data["Fare"].fillna(data["Fare"].median())

    data["FamilySize"] = data["SibSp"] + data["Parch"] + 1

    target = data["Survived"]
    features = data[
        ["Pclass", "Sex", "Age", "SibSp", "Parch", "Fare", "Embarked", "FamilySize"]
    ].copy()

    return features, target


def build_preprocessor() -> ColumnTransformer:
    """Create the column transformer for numeric and categorical feature pipelines."""
    numerical_features = ["Age", "Fare", "SibSp", "Parch", "FamilySize"]
    categorical_features = ["Pclass", "Sex", "Embarked"]

    numerical_pipeline = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="median")),
            ("scaler", StandardScaler()),
        ]
    )

    categorical_pipeline = Pipeline(
        steps=[
            ("imputer", SimpleImputer(strategy="most_frequent")),
            ("encoder", OneHotEncoder(drop="first", handle_unknown="ignore")),
        ]
    )

    preprocessor = ColumnTransformer(
        transformers=[
            ("num", numerical_pipeline, numerical_features),
            ("cat", categorical_pipeline, categorical_features),
        ],
        remainder="drop",
    )

    return preprocessor


def train_test_split_data(features: pd.DataFrame, target: pd.Series) -> tuple:
    """Create the stratified train/test split for model development."""
    X_train, X_test, y_train, y_test = train_test_split(
        features,
        target,
        test_size=0.20,
        random_state=42,
        stratify=target,
    )
    return X_train, X_test, y_train, y_test


# -----------------------------------------------------------------------------
# Model definitions and evaluation
# -----------------------------------------------------------------------------
def make_models() -> dict[str, object]:
    """Return the two model objects required by the assignment."""
    return {
        "Logistic Regression": LogisticRegression(random_state=42, max_iter=1000),
        "Random Forest": RandomForestClassifier(
            n_estimators=100,
            max_depth=5,
            random_state=42,
        ),
    }


def evaluate_model(model, X_train, X_test, y_train, y_test) -> dict:
    """Train a model and compute the required evaluation metrics."""
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_proba = model.predict_proba(X_test)[:, 1] if hasattr(model, "predict_proba") else None

    metrics = {
        "accuracy": accuracy_score(y_test, y_pred),
        "precision": precision_score(y_test, y_pred, zero_division=0),
        "recall": recall_score(y_test, y_pred, zero_division=0),
        "f1": f1_score(y_test, y_pred, zero_division=0),
        "roc_auc": roc_auc_score(y_test, y_proba) if y_proba is not None else np.nan,
        "confusion_matrix": confusion_matrix(y_test, y_pred),
        "y_pred": y_pred,
        "y_proba": y_proba,
    }

    return metrics


def create_confusion_matrix_figure(cm_1: np.ndarray, cm_2: np.ndarray) -> Path:
    """Create a dual-panel confusion matrix figure for the two models."""
    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    labels = ["Died", "Survived"]
    for ax, cm, title in zip(
        axes,
        [cm_1, cm_2],
        ["Logistic Regression", "Random Forest"],
    ):
        sns.heatmap(
            cm,
            annot=True,
            fmt="d",
            cmap="Blues",
            square=True,
            cbar=False,
            xticklabels=labels,
            yticklabels=labels,
            ax=ax,
        )
        ax.set_title(title, fontsize=12, weight="bold")
        ax.set_xlabel("Predicted")
        ax.set_ylabel("Actual")

    fig.tight_layout()
    path = OUTPUT_DIR / "fig1_confusion_matrices.png"
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


def create_roc_figure(model_results: dict[str, dict], X_test: pd.DataFrame, y_test: pd.Series) -> Path:
    """Create ROC curves for both models with the random-chance benchmark line."""
    fig, ax = plt.subplots(figsize=(9, 6))

    ax.plot([0, 1], [0, 1], linestyle="--", color="gray", linewidth=1.5, label="Random Chance")

    for name, result in model_results.items():
        if result.get("y_proba") is None:
            continue
        fpr, tpr, _ = roc_curve(y_test, result["y_proba"])
        ax.plot(fpr, tpr, linewidth=2, label=f"{name} (AUC = {result['roc_auc']:.3f})")

    ax.set_title("ROC Curves for Titanic Survival Models", fontsize=14, weight="bold")
    ax.set_xlabel("False Positive Rate")
    ax.set_ylabel("True Positive Rate")
    ax.legend(loc="lower right")
    ax.grid(alpha=0.2)
    fig.tight_layout()

    path = OUTPUT_DIR / "fig2_roc_curves.png"
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return path


# -----------------------------------------------------------------------------
# Word document generation
# -----------------------------------------------------------------------------
def add_centered_image(doc: Document, image_path: Path, width_in_inches: float = 5.5) -> None:
    """Insert a centered image into the report."""
    doc.add_picture(str(image_path), width=Inches(width_in_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def format_performance_table(results: dict[str, dict]) -> pd.DataFrame:
    """Build a compact performance comparison table for the report."""
    rows = []
    for model_name, model_metrics in results.items():
        rows.append(
            {
                "Model": model_name,
                "Accuracy": round(model_metrics["accuracy"], 4),
                "Precision": round(model_metrics["precision"], 4),
                "Recall": round(model_metrics["recall"], 4),
                "F1-Score": round(model_metrics["f1"], 4),
                "ROC-AUC": round(model_metrics["roc_auc"], 4),
            }
        )
    return pd.DataFrame(rows)


def build_document(results: dict[str, dict], figure_paths: dict[str, Path]) -> Path:
    """Create the final Word report with narrative and figures."""
    doc = Document()
    doc.add_heading("Week 4: Machine Learning Model Development & Evaluation", 0)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This project evaluated two supervised learning models for predicting Titanic survival using a compact, interpretable feature set. The workflow included feature engineering, imputation, preprocessing, stratified train/test splitting, model training, classification evaluation, and visual analysis. Logistic regression provides a strong baseline with transparent coefficients, while the random forest model offers non-linear structure capture and improved flexibility. In practical terms, the goal was to assess not only overall predictive performance but also which model balances accuracy, recall, and calibration trade-offs for a binary classification problem with meaningful operational implications."
    )

    doc.add_heading("Methodological Description", level=1)
    doc.add_paragraph(
        "The Titanic dataset was downloaded and preprocessed using a structured feature engineering workflow. Family size was derived as SibSp + Parch + 1 to capture household structure. A ColumnTransformer was then built with separate pipelines for numerical and categorical data: median imputation and StandardScaler for Age, Fare, SibSp, Parch, and FamilySize, and mode imputation plus OneHotEncoder(drop='first') for Pclass, Sex, and Embarked. The target variable was survival status, and an 80/20 stratified train-test split with random_state=42 was used to preserve the class balance across training and evaluation sets."
    )

    doc.add_heading("Quantitative Performance Data", level=1)
    performance_table = format_performance_table(results)
    table = doc.add_table(rows=1, cols=len(performance_table.columns))
    table.style = "Light List Accent 1"

    headers = list(performance_table.columns)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)

    for _, row in performance_table.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    # Add a brief sentence summarizing performance.
    best_model = max(results, key=lambda name: results[name]["f1"])
    best_auc = max(results[name]["roc_auc"] for name in results)
    doc.add_paragraph(
        f"The best-performing model by F1-score was {best_model}. The highest ROC-AUC observed was {best_auc:.3f}, indicating the strongest discriminatory power across the two models tested."
    )

    doc.add_heading("Figure 1: Confusion Matrices", level=1)
    add_centered_image(doc, figure_paths["fig1_confusion_matrices"], width_in_inches=5.5)
    doc.add_paragraph(
        "Figure 1 compares the classification error patterns of the logistic regression and random forest models. A confusion matrix highlights how often each model correctly identifies survivors and non-survivors and where it makes false positive or false negative mistakes. This comparison is useful because a model with strong overall accuracy can still perform poorly on one class if the error distribution is imbalanced."
    )

    doc.add_heading("Figure 2: ROC Curves", level=1)
    add_centered_image(doc, figure_paths["fig2_roc_curves"], width_in_inches=5.5)
    doc.add_paragraph(
        "Figure 2 overlays the ROC curves for both models and includes a random-chance benchmark. The ROC curve summarises the trade-off between sensitivity and specificity over a range of decision thresholds, while the area under the curve provides a threshold-independent measure of discriminative power. This is particularly relevant when the operational cost of false negatives differs from false positives."
    )

    doc.add_heading("Critical Discussion", level=1)
    doc.add_paragraph(
        "Potential sources of error include the omission of variables not present in the dataset, the possibility that historical data reflect social biases, and the fact that the train/test split may not capture all temporal or population-level variation. The logistic regression model is interpretable but may underfit if the relationship between predictors and survival is non-linear. Conversely, the random forest model can capture non-linear interactions but may overfit if the tree depth or minimum split criteria are not tuned carefully."
    )
    doc.add_paragraph(
        "Future improvements should include a more systematic hyperparameter search using GridSearchCV, model comparison against stronger baselines such as XGBoost, and validation using cross-validation to better estimate robustness. Additional diagnostics such as calibration curves, precision-recall curves, and feature importance analysis would support more nuanced model selection and operational deployment decisions."
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


# -----------------------------------------------------------------------------
# Main pipeline execution
# -----------------------------------------------------------------------------
def main() -> None:
    """Execute the full machine-learning workflow and generate the final report."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    raw_df = fetch_data()
    features, target = prepare_dataset(raw_df)
    X_train, X_test, y_train, y_test = train_test_split_data(features, target)

    preprocessor = build_preprocessor()
    X_train_processed = preprocessor.fit_transform(X_train)
    X_test_processed = preprocessor.transform(X_test)

    models = make_models()
    model_results = {}

    for name, model in models.items():
        # Rebuild a fresh pipeline for each model with the same preprocessing logic.
        pipeline = Pipeline(
            steps=[
                ("preprocessor", preprocessor),
                ("model", model),
            ]
        )
        pipeline.fit(X_train, y_train)
        y_pred = pipeline.predict(X_test)
        y_proba = pipeline.predict_proba(X_test)[:, 1]

        model_results[name] = {
            "accuracy": accuracy_score(y_test, y_pred),
            "precision": precision_score(y_test, y_pred, zero_division=0),
            "recall": recall_score(y_test, y_pred, zero_division=0),
            "f1": f1_score(y_test, y_pred, zero_division=0),
            "roc_auc": roc_auc_score(y_test, y_proba),
            "confusion_matrix": confusion_matrix(y_test, y_pred),
            "y_proba": y_proba,
        }

    # Use the model results to generate the required figures.
    logistic_cm = model_results["Logistic Regression"]["confusion_matrix"]
    rf_cm = model_results["Random Forest"]["confusion_matrix"]
    fig1_path = create_confusion_matrix_figure(logistic_cm, rf_cm)
    fig2_path = create_roc_figure(model_results, X_test, y_test)

    report_path = build_document(
        model_results,
        {
            "fig1_confusion_matrices": fig1_path,
            "fig2_roc_curves": fig2_path,
        },
    )

    LOGGER.info("Saved ML outputs to: %s", OUTPUT_DIR)
    LOGGER.info("Saved report to: %s", report_path)
    print(f"Report created successfully: {report_path}")


if __name__ == "__main__":
    main()
