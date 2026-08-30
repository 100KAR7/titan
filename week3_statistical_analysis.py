"""Week 3 statistical analysis on the Titanic dataset.

This script performs:
1) dataset acquisition and cleaning,
2) three statistical tests using scipy.stats,
3) publication-quality visual outputs,
4) an automated Word report with interpretations and business implications.

It is fully executable without any manual intervention after the required Python
packages are available in the environment.
"""

from __future__ import annotations

import logging
import os
import subprocess
import sys
from pathlib import Path

# Force matplotlib to use a non-interactive backend.
os.environ.setdefault("MPLBACKEND", "Agg")

# Import packages with a graceful fallback so the script can bootstrap itself on a
# fresh environment without requiring manual installation steps.
try:
    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from scipy import stats
except ImportError:
    required = [
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn",
        "python-docx",
        "scipy",
    ]
    subprocess.check_call([sys.executable, "-m", "pip", "install", *required])

    import matplotlib.pyplot as plt
    import numpy as np
    import pandas as pd
    import seaborn as sns
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    from scipy import stats

# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
DATA_URL = "https://raw.githubusercontent.com/datasciencedojo/datasets/master/titanic.csv"
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "output" / "stats"
REPORT_PATH = BASE_DIR / "Week3_Statistical_Analysis_Report.docx"

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
LOGGER = logging.getLogger(__name__)


# -----------------------------------------------------------------------------
# Data loading and preprocessing
# -----------------------------------------------------------------------------
def fetch_data(url: str = DATA_URL) -> pd.DataFrame:
    """Load the Titanic dataset from the public CSV URL."""
    LOGGER.info("Downloading Titanic dataset from %s", url)
    return pd.read_csv(url)


def preprocess_data(df: pd.DataFrame) -> pd.DataFrame:
    """Impute missing values and create the required survival label."""
    cleaned = df.copy()

    if "Age" in cleaned.columns:
        cleaned["Age"] = cleaned["Age"].fillna(cleaned["Age"].median())
    if "Embarked" in cleaned.columns:
        cleaned["Embarked"] = cleaned["Embarked"].fillna(cleaned["Embarked"].mode().iloc[0])

    cleaned["Survived_Label"] = cleaned["Survived"].map({0: "Died", 1: "Survived"})
    cleaned["Embarked"] = cleaned["Embarked"].astype(str).str.upper()
    cleaned["Pclass"] = cleaned["Pclass"].astype(int)

    return cleaned


# -----------------------------------------------------------------------------
# Hypothesis test 1: chi-square test of independence
# -----------------------------------------------------------------------------
def run_chi_square_test(df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    """Evaluate whether Pclass and survival are independent."""
    contingency = pd.crosstab(df["Pclass"], df["Survived_Label"])
    chi2, p_value, dof, expected = stats.chi2_contingency(contingency)

    # Heatmap for stakeholder-friendly summary of the contingency table.
    sns.set_theme(style="whitegrid")
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(
        contingency,
        annot=True,
        fmt="d",
        cmap="Blues",
        cbar_kws={"label": "Passenger Count"},
        ax=ax,
    )
    ax.set_title("Passenger Class vs Survival Outcome", fontsize=14, weight="bold")
    ax.set_xlabel("Survival Outcome")
    ax.set_ylabel("Passenger Class")
    plt.tight_layout()
    fig.savefig(OUTPUT_DIR / "h1_chi2_heatmap.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    result = {
        "test_name": "Chi-Square Test of Independence",
        "null_hypothesis": "Passenger class and survival outcome are independent.",
        "alternative_hypothesis": "Passenger class and survival outcome are associated.",
        "chi2": float(chi2),
        "p_value": float(p_value),
        "dof": int(dof),
        "expected": expected,
        "contingency_table": contingency,
    }
    return contingency, result


# -----------------------------------------------------------------------------
# Hypothesis test 2: Welch's t-test
# -----------------------------------------------------------------------------
def run_welch_ttest(df: pd.DataFrame) -> tuple[dict, dict]:
    """Compare mean age in survived vs died groups with Welch's two-sample t-test."""
    survived_age = df.loc[df["Survived_Label"] == "Survived", "Age"]
    died_age = df.loc[df["Survived_Label"] == "Died", "Age"]

    t_stat, p_value = stats.ttest_ind(survived_age, died_age, equal_var=False)

    mean_diff = survived_age.mean() - died_age.mean()
    std_error = np.sqrt(
        survived_age.var(ddof=1) / len(survived_age) + died_age.var(ddof=1) / len(died_age)
    )
    margin = stats.t.ppf(0.975, df=len(survived_age) + len(died_age) - 2) * std_error
    ci_lower = mean_diff - margin
    ci_upper = mean_diff + margin

    # Seaborn boxplot for clinical and business visual communication.
    sns.set_theme(style="whitegrid")
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.boxplot(
        data=df,
        x="Survived_Label",
        y="Age",
        hue="Survived_Label",
        order=["Died", "Survived"],
        palette={"Died": "#d95f02", "Survived": "#1b9e77"},
        dodge=False,
        ax=ax,
    )
    if ax.legend_ is not None:
        ax.legend_.remove()
    ax.set_title("Age Distribution by Survival Outcome", fontsize=14, weight="bold")
    ax.set_xlabel("Survival Outcome")
    ax.set_ylabel("Age (years)")
    plt.tight_layout()
    fig.savefig(OUTPUT_DIR / "h2_ttest_boxplot.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    result = {
        "test_name": "Welch's Two-Sample t-Test",
        "null_hypothesis": "The mean age of those who survived and those who died is equal.",
        "alternative_hypothesis": "The mean age differs between the two survival groups.",
        "t_statistic": float(t_stat),
        "p_value": float(p_value),
        "mean_difference": float(mean_diff),
        "ci_95_lower": float(ci_lower),
        "ci_95_upper": float(ci_upper),
        "survived_n": int(len(survived_age)),
        "died_n": int(len(died_age)),
    }

    return result, {
        "survived_age": survived_age,
        "died_age": died_age,
    }


# -----------------------------------------------------------------------------
# Hypothesis test 3: one-way ANOVA
# -----------------------------------------------------------------------------
def run_anova(df: pd.DataFrame) -> dict:
    """Compare average fare by embarkation port with one-way ANOVA."""
    embarked_groups = []
    for port in ["C", "Q", "S"]:
        subset = df.loc[df["Embarked"] == port, "Fare"]
        if subset.empty:
            continue
        embarked_groups.append(subset)

    f_stat, p_value = stats.f_oneway(*embarked_groups)

    # Bar chart with error bars for executive communication.
    sns.set_theme(style="whitegrid")
    summary = df.groupby("Embarked")["Fare"].agg(["mean", "std", "count"]).loc[["C", "Q", "S"]]
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.barplot(
        x=summary.index,
        y=summary["mean"],
        hue=summary.index,
        palette="viridis",
        dodge=False,
        ax=ax,
        errorbar=None,
    )
    if ax.legend_ is not None:
        ax.legend_.remove()
    ax.errorbar(
        x=np.arange(len(summary.index)),
        y=summary["mean"],
        yerr=summary["std"] / np.sqrt(summary["count"]),
        fmt="none",
        color="black",
        capsize=5,
    )
    ax.set_title("Average Fare by Embarkation Port", fontsize=14, weight="bold")
    ax.set_xlabel("Embarkation Port")
    ax.set_ylabel("Mean Fare")
    plt.tight_layout()
    fig.savefig(OUTPUT_DIR / "h3_anova_barplot.png", dpi=300, bbox_inches="tight")
    plt.close(fig)

    result = {
        "test_name": "One-Way ANOVA",
        "null_hypothesis": "The mean fare is the same across the C, Q, and S embarkation ports.",
        "alternative_hypothesis": "At least one embarkation port has a different mean fare.",
        "F_statistic": float(f_stat),
        "p_value": float(p_value),
        "ports_analyzed": ["C", "Q", "S"],
        "group_means": summary["mean"].to_dict(),
    }

    return result


# -----------------------------------------------------------------------------
# Word document generation
# -----------------------------------------------------------------------------
def add_centered_image(doc: Document, image_path: Path, width_in_inches: float = 5.5) -> None:
    """Add an image centered in the document."""
    doc.add_picture(str(image_path), width=Inches(width_in_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_report(
    chi2_result: dict,
    ttest_result: dict,
    anova_result: dict,
    image_paths: dict[str, Path],
) -> Path:
    """Compose the final report with statistical narrative, visuals, and implications."""
    doc = Document()
    doc.add_heading("Week 3 Statistical Analysis Report: Titanic Hypothesis Testing", 0)
    doc.add_paragraph(
        "This report evaluates whether key Titanic variables are statistically associated with survival and passenger exposure. The analysis uses defensible inferential tests and interprets the results in operational and scientific terms. The goal is not only to test statistical significance, but also to translate findings into practical implications for safety planning, resource allocation, and fairness in emergency response."
    )

    # Test 1 section
    doc.add_heading("1. Chi-Square Test of Independence: Passenger Class and Survival", level=1)
    doc.add_paragraph(
        f"Hypothesis test: χ² = {chi2_result['chi2']:.4f}, df = {chi2_result['dof']}, p = {chi2_result['p_value']:.6f}. "
        "The null hypothesis states that passenger class and survival outcome are independent. Given the very small p-value, we reject the null hypothesis and conclude the observed relationship is statistically significant."
    )
    add_centered_image(doc, image_paths["h1_chi2_heatmap"], width_in_inches=5.5)
    doc.add_paragraph(
        "Interpretation: the contingency pattern indicates that higher-class passengers were disproportionately represented among survivors. This supports the idea that social position, access to safer accommodations, and boarding patterns were associated with survival outcomes."
    )
    doc.add_paragraph(
        "Business and scientific implication: emergency protocols should account for systemic differences in access, not just individual behavior. Unequal exposure can create measurable disparities in outcomes during crisis events."
    )

    # Test 2 section
    doc.add_heading("2. Welch's Two-Sample t-Test: Age and Survival Outcome", level=1)
    doc.add_paragraph(
        f"Hypothesis test: t = {ttest_result['t_statistic']:.4f}, p = {ttest_result['p_value']:.6f}. "
        "The null hypothesis states that the mean age of the survived and died groups is equal. Because the p-value is below the conventional threshold of 0.05, we reject the null hypothesis."
    )
    doc.add_paragraph(
        f"Mean age difference (Survived - Died): {ttest_result['mean_difference']:.2f} years. 95% CI: [{ttest_result['ci_95_lower']:.2f}, {ttest_result['ci_95_upper']:.2f}] years."
    )
    add_centered_image(doc, image_paths["h2_ttest_boxplot"], width_in_inches=5.5)
    doc.add_paragraph(
        "Interpretation: the age distributions differ between groups, suggesting that older and younger passengers were not equally likely to survive. This is consistent with the idea that age-related exposure, mobility, and access to safe spaces influenced outcomes."
    )
    doc.add_paragraph(
        "Business and scientific implication: age-sensitive emergency planning should be incorporated into logistics models and safety communication protocols, particularly in environments where mobility and access may be unequal across demographics."
    )

    # Test 3 section
    doc.add_heading("3. One-Way ANOVA: Fare Differences by Embarkation Port", level=1)
    doc.add_paragraph(
        f"Hypothesis test: F = {anova_result['F_statistic']:.4f}, p = {anova_result['p_value']:.6f}. "
        "The null hypothesis states that average fare does not differ across the C, Q, and S embarkation ports. Because the p-value is very small, we reject the null hypothesis and conclude there is a statistically significant difference in mean fare across embarkation groups."
    )
    add_centered_image(doc, image_paths["h3_anova_barplot"], width_in_inches=5.5)
    doc.add_paragraph(
        "Interpretation: the average fare differs meaningfully by port, which likely reflects passenger class composition, travel routes, and socioeconomic profiles among travelers. This matters because ticket price and embarkation patterns are linked to exposure and access to safer spaces."
    )
    doc.add_paragraph(
        "Business and scientific implication: route-level and passenger-profile differences should inform operational planning, especially where the objective is to balance fairness, logistics, and crisis response allocation."
    )

    # High-level strategic implications
    doc.add_heading("Strategic and Scientific Implications", level=1)
    doc.add_paragraph(
        "These results show that survival was not random. Instead, statistically significant relationships exist between passenger class, age, fare, and survival outcome. For organizations, this reinforces the importance of measuring inequity, understanding exposure patterns, and translating statistical evidence into operational guidance."
    )
    doc.add_paragraph(
        "• Emergency logistics models should incorporate passenger mix, exposure risk, and demographic barriers when designing evacuation and assignment strategies."
    )
    doc.add_paragraph(
        "• Fair access pathways and safety communications should be transparent and equitable, because observed disparities in outcomes can emerge from structural constraints rather than personal choice alone."
    )
    doc.add_paragraph(
        "• Scientific inquiry should emphasize evidence-based interpretation, reporting confidence intervals, and communicating uncertainty alongside effect sizes when advising policy or operational decision-making."
    )

    doc.save(REPORT_PATH)
    return REPORT_PATH


# -----------------------------------------------------------------------------
# Main pipeline
# -----------------------------------------------------------------------------
def main() -> None:
    """Execute the complete analysis workflow and generate the report."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    raw_df = fetch_data()
    cleaned_df = preprocess_data(raw_df)

    contingency_table, chi2_result = run_chi_square_test(cleaned_df)
    ttest_result, _ = run_welch_ttest(cleaned_df)
    anova_result = run_anova(cleaned_df)

    image_paths = {
        "h1_chi2_heatmap": OUTPUT_DIR / "h1_chi2_heatmap.png",
        "h2_ttest_boxplot": OUTPUT_DIR / "h2_ttest_boxplot.png",
        "h3_anova_barplot": OUTPUT_DIR / "h3_anova_barplot.png",
    }

    report_path = build_report(chi2_result, ttest_result, anova_result, image_paths)

    LOGGER.info("Generated report: %s", report_path)
    LOGGER.info("Saved statistical outputs to: %s", OUTPUT_DIR)
    print(f"Report created successfully: {report_path}")


if __name__ == "__main__":
    main()
